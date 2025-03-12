from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.document_loaders import DirectoryLoader, UnstructuredFileLoader
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
import os
import shutil
from dotenv import load_dotenv
from typing import List, Dict, Any, Optional
import fitz  # PyMuPDF for PDF processing
from docx import Document as DocxDocument  # python-docx for DOCX processing
import re
import hashlib
from tqdm import tqdm
import concurrent.futures
import logging

# Set up logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                    handlers=[logging.FileHandler("document_processing.log"),
                              logging.StreamHandler()])
logger = logging.getLogger(__name__)

load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")
CHROMA_PATH = "chroma"
DATA_PATH = "documents"
CACHE_DIR = "document_cache"

# Ensure cache directory exists
os.makedirs(CACHE_DIR, exist_ok=True)

def generate_data_store():
    """Main function to process documents and create the vector store."""
    logger.info("Starting document processing pipeline")
    
    # Load documents with progress tracking
    with tqdm(desc="Loading Documents") as pbar:
        documents = load_documents(DATA_PATH, progress_callback=lambda: pbar.update(1))
    
    logger.info(f"Loaded {len(documents)} documents.")
    
    # Extract metadata from documents
    with tqdm(total=len(documents), desc="Extracting Metadata") as pbar:
        documents = extract_document_metadata(documents, progress_callback=lambda: pbar.update(1))
    
    # Split documents into chunks
    chunks = split_text(documents)
    logger.info(f"Generated {len(chunks)} chunks.")
    
    # Save to Chroma DB
    save_to_chroma(chunks)
    logger.info("Document processing complete.")

def generate_document_hash(file_path: str) -> str:
    """Generate a hash of the document content for caching purposes."""
    with open(file_path, 'rb') as f:
        file_content = f.read()
        return hashlib.sha256(file_content).hexdigest()

def load_from_cache(file_path: str) -> Optional[str]:
    """Load document text from cache if available."""
    doc_hash = generate_document_hash(file_path)
    cache_file = os.path.join(CACHE_DIR, f"{doc_hash}.txt")
    
    if os.path.exists(cache_file):
        with open(cache_file, 'r', encoding='utf-8') as f:
            return f.read()
    return None

def save_to_cache(file_path: str, text: str):
    """Save document text to cache."""
    doc_hash = generate_document_hash(file_path)
    cache_file = os.path.join(CACHE_DIR, f"{doc_hash}.txt")
    
    with open(cache_file, 'w', encoding='utf-8') as f:
        f.write(text)

def process_single_document(file_path: str) -> Optional[Document]:
    """Process a single document and return a Document object."""
    try:
        # Check cache first
        cached_text = load_from_cache(file_path)
        if cached_text:
            logger.info(f"Loading {file_path} from cache")
            return Document(page_content=cached_text, metadata={"source": file_path})
        
        # Process based on file extension
        if file_path.endswith('.pdf'):
            text = extract_pdf_text(file_path)
        elif file_path.endswith('.docx'):
            text = extract_docx_text(file_path)
        elif file_path.endswith('.doc'):
            # For .doc files, use UnstructuredFileLoader as fallback
            loader = UnstructuredFileLoader(file_path)
            doc = loader.load()[0]
            text = doc.page_content
        else:
            logger.warning(f"Unsupported file type: {file_path}")
            return None
            
        # Save to cache
        save_to_cache(file_path, text)
        
        return Document(page_content=text, metadata={"source": file_path})
    except Exception as e:
        logger.error(f"Error processing {file_path}: {str(e)}")
        return None

def load_documents(data_path: str, progress_callback=None) -> List[Document]:
    """Load documents from various file formats with parallel processing."""
    supported_extensions = ['.pdf', '.docx', '.doc']
    
    # Get all files with supported extensions
    all_files = []
    for root, _, files in os.walk(data_path):
        for file in files:
            if any(file.endswith(ext) for ext in supported_extensions):
                all_files.append(os.path.join(root, file))
    
    documents = []
    
    # Process documents in parallel
    with concurrent.futures.ThreadPoolExecutor(max_workers=os.cpu_count()) as executor:
        future_to_file = {executor.submit(process_single_document, file): file for file in all_files}
        
        for future in concurrent.futures.as_completed(future_to_file):
            file = future_to_file[future]
            try:
                doc = future.result()
                if doc:
                    documents.append(doc)
                    if progress_callback:
                        progress_callback()
            except Exception as e:
                logger.error(f"Error processing {file}: {str(e)}")
    
    return documents

def extract_pdf_text(pdf_path: str) -> str:
    """Extract text from a PDF with advanced processing."""
    doc = fitz.open(pdf_path)
    text_parts = []
    
    for page_num, page in enumerate(doc):
        # Extract text with formatting information
        text = page.get_text("text")
        
        # Add page number metadata directly in the text
        page_info = f"\n--- Page {page_num + 1} ---\n"
        text_parts.append(page_info + text)
    
    return "\n".join(text_parts)

def extract_docx_text(docx_path: str) -> str:
    """Extract text from a DOCX file with structure preservation."""
    doc = DocxDocument(docx_path)
    text_parts = []
    
    # Extract headings and paragraphs
    for para in doc.paragraphs:
        # Check if it's a heading
        if para.style.name.startswith('Heading'):
            heading_level = para.style.name.replace('Heading', '').strip()
            if heading_level:
                # Format headings appropriately
                text_parts.append(f"\n{'#' * int(heading_level)} {para.text}")
            else:
                text_parts.append(f"\n# {para.text}")
        else:
            text_parts.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            table_text.append(" | ".join(row_text))
        text_parts.append("\n" + "\n".join(table_text) + "\n")
    
    return "\n".join(text_parts)

def extract_document_metadata(documents: List[Document], progress_callback=None) -> List[Document]:
    """Extract and add enhanced metadata to documents."""
    enhanced_docs = []
    
    for doc in documents:
        file_path = doc.metadata.get('source', '')
        filename = os.path.basename(file_path)
        file_extension = os.path.splitext(filename)[1].lower()
        
        # Common metadata for all documents
        metadata = {
            "source": file_path,
            "filename": filename,
            "extension": file_extension,
            "relative_path": os.path.relpath(file_path, start=DATA_PATH),
            "file_size_bytes": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
        }
        
        # Extract document-specific metadata
        if file_extension == '.pdf':
            try:
                pdf_metadata = extract_pdf_metadata(file_path)
                metadata.update(pdf_metadata)
            except Exception as e:
                logger.error(f"Error extracting PDF metadata from {file_path}: {str(e)}")
        
        elif file_extension == '.docx':
            try:
                docx_metadata = extract_docx_metadata(file_path)
                metadata.update(docx_metadata)
            except Exception as e:
                logger.error(f"Error extracting DOCX metadata from {file_path}: {str(e)}")
        
        # Add document type classification
        metadata["document_type"] = classify_document_type(doc.page_content, filename)
        
        # Create new document with enhanced metadata
        enhanced_doc = Document(page_content=doc.page_content, metadata=metadata)
        enhanced_docs.append(enhanced_doc)
        
        if progress_callback:
            progress_callback()
    
    return enhanced_docs

def extract_pdf_metadata(pdf_path: str) -> Dict[str, Any]:
    """Extract metadata from PDF files."""
    doc = fitz.open(pdf_path)
    metadata = {
        "page_count": doc.page_count,
        "pdf_title": doc.metadata.get('title', ''),
        "pdf_author": doc.metadata.get('author', ''),
        "pdf_subject": doc.metadata.get('subject', ''),
        "pdf_creator": doc.metadata.get('creator', ''),
        "pdf_producer": doc.metadata.get('producer', ''),
    }
    doc.close()
    return metadata

def extract_docx_metadata(docx_path: str) -> Dict[str, Any]:
    """Extract metadata from DOCX files."""
    doc = DocxDocument(docx_path)
    core_props = doc.core_properties
    
    # Count headings by level
    heading_counts = {}
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            heading_level = para.style.name.replace('Heading', '').strip()
            if heading_level:
                level = int(heading_level)
                heading_counts[f"heading{level}_count"] = heading_counts.get(f"heading{level}_count", 0) + 1
    
    return {
        "docx_title": core_props.title or '',
        "docx_author": core_props.author or '',
        "docx_subject": core_props.subject or '',
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        **heading_counts
    }

def classify_document_type(content: str, filename: str) -> str:
    """
    Classify document into categories based on content and filename patterns.
    """
    filename_lower = filename.lower()
    content_sample = content[:5000].lower()  # Use the first 5000 chars for classification
    
    # Research paper indicators
    paper_patterns = [
        r'\babstract\b', r'\bintroduction\b', r'\bmethodology\b', r'\bresults\b', 
        r'\bdiscussion\b', r'\bconclusion\b', r'\breferences\b', r'\bcitation\b',
        r'\bliterature review\b'
    ]
    
    if any(re.search(pattern, content_sample) for pattern in paper_patterns):
        return "research_paper"
    
    # Report patterns
    report_patterns = [
        r'\breport\b', r'\bfindings\b', r'\bexecutive summary\b', 
        r'\bquarterly\b', r'\bannual\b'
    ]
    
    if any(re.search(pattern, content_sample) for pattern in report_patterns):
        return "report"
    
    # Manual/guide patterns
    guide_patterns = [
        r'\bmanual\b', r'\bguide\b', r'\binstructions\b', r'\btutorial\b',
        r'\bstep[-\s]by[-\s]step\b', r'\bhow[-\s]to\b'
    ]
    
    if any(re.search(pattern, content_sample) for pattern in guide_patterns):
        return "guide"
        
    # Default to generic document
    return "document"

def split_text(documents: List[Document]) -> List[Document]:
    """Split documents into smaller chunks with advanced configuration."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        add_start_index=True,
        strip_whitespace=True,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    
    logger.info(f"Splitting {len(documents)} documents into chunks...")
    chunks = text_splitter.split_documents(documents)
    
    # Add chunk metadata
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = i
        
        # Extract page information if present in the content
        page_match = re.search(r"--- Page (\d+) ---", chunk.page_content)
        if page_match:
            chunk.metadata["page"] = int(page_match.group(1))
            # Optionally remove the page marker from the content
            chunk.page_content = re.sub(r"--- Page \d+ ---\n?", "", chunk.page_content)
    
    if chunks:
        logger.info(f"First chunk sample: {chunks[0].page_content[:100]}...")
        logger.info(f"First chunk metadata: {chunks[0].metadata}")
    
    return chunks

def save_to_chroma(chunks: List[Document]):
    """Save document chunks to Chroma vector store with optimized settings."""
    if os.path.exists(CHROMA_PATH):
        logger.info(f"Removing existing vector store at {CHROMA_PATH}")
        shutil.rmtree(CHROMA_PATH)
    
    logger.info(f"Creating new vector store with {len(chunks)} chunks...")
    
    # Use a more advanced embedding model
    embedding_function = OpenAIEmbeddings(
        openai_api_key=api_key,
        model="text-embedding-3-large",  # Using the more advanced model
        dimensions=1536,  # Explicitly setting dimensions
    )
    
    # Create and persist the database with additional settings
    db = Chroma.from_documents(
        chunks,
        embedding_function,
        persist_directory=CHROMA_PATH,
        collection_metadata={"hnsw:space": "cosine"}  # Explicitly set distance metric
    )
    
    db.persist()
    logger.info(f"Successfully saved {len(chunks)} chunks to {CHROMA_PATH}.")